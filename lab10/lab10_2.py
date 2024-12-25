from docx import Document
from docx.shared import Inches

document = Document('lab10/document.docx')

document.add_picture('lab10/screen.png', width=Inches(4.0))
paragraph = document.add_paragraph()
paragraph.add_run('Текстовая подпись под изображением')
document.save('lab10/document.docx')